from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Literal


BlockKind = Literal["paragraph", "bullets", "numbered_list", "table", "code"]


@dataclass
class DocumentBlock:
    kind: BlockKind
    text: str = ""
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class DocumentSection:
    heading: str
    level: int = 1
    blocks: list[DocumentBlock] = field(default_factory=list)


@dataclass
class DocumentAST:
    title: str
    sections: list[DocumentSection] = field(default_factory=list)
    metadata: dict[str, str] = field(default_factory=dict)


@dataclass
class DocumentQaIssue:
    code: str
    message: str


def document_ast_from_markdown(title: str, markdown: str) -> DocumentAST:
    """Parse the v3 markdown draft into a small Word-oriented document AST."""
    ast = DocumentAST(title=_clean_text(title) or "Fronei document")
    current = DocumentSection(heading="Overview", level=1)
    paragraph_lines: list[str] = []
    bullet_items: list[str] = []
    numbered_items: list[str] = []
    code_lines: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        text = _clean_text(" ".join(paragraph_lines))
        if text:
            current.blocks.append(DocumentBlock(kind="paragraph", text=text))
        paragraph_lines = []

    def flush_bullets() -> None:
        nonlocal bullet_items
        if bullet_items:
            current.blocks.append(DocumentBlock(kind="bullets", items=bullet_items))
        bullet_items = []

    def flush_numbered() -> None:
        nonlocal numbered_items
        if numbered_items:
            current.blocks.append(DocumentBlock(kind="numbered_list", items=numbered_items))
        numbered_items = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            current.blocks.append(DocumentBlock(kind="code", text="\n".join(code_lines).strip()))
        code_lines = []

    def flush_all() -> None:
        flush_paragraph()
        flush_bullets()
        flush_numbered()
        flush_code()

    def finish_section() -> None:
        flush_all()
        if current.blocks or current.heading != "Overview":
            ast.sections.append(current)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                flush_bullets()
                flush_numbered()
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            finish_section()
            heading_text = _clean_heading(heading.group(2))
            level = max(1, min(3, len(heading.group(1))))
            if not ast.sections and level == 1 and _same_heading(heading_text, ast.title):
                current = DocumentSection(heading="Overview", level=1)
            else:
                current = DocumentSection(heading=heading_text, level=level)
            index += 1
            continue

        if _is_table_start(lines, index):
            flush_paragraph()
            flush_bullets()
            flush_numbered()
            table_rows: list[list[str]] = []
            while index < len(lines) and "|" in lines[index]:
                candidate = lines[index].strip()
                if _is_table_separator(candidate):
                    index += 1
                    continue
                row = [_clean_inline(cell.strip()) for cell in candidate.strip("|").split("|")]
                if any(row):
                    table_rows.append(row)
                index += 1
            normalized = _normalize_table_rows(table_rows)
            if normalized:
                current.blocks.append(DocumentBlock(kind="table", rows=normalized))
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            flush_numbered()
            bullet_items.append(_clean_inline(bullet.group(1)))
            index += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            flush_bullets()
            numbered_items.append(_clean_inline(numbered.group(1)))
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            flush_bullets()
            flush_numbered()
            index += 1
            continue

        flush_bullets()
        flush_numbered()
        paragraph_lines.append(_clean_inline(stripped))
        index += 1

    finish_section()
    ast.sections = [section for section in ast.sections if section.blocks or section.heading]
    if not ast.sections:
        ast.sections.append(DocumentSection(heading="Overview", blocks=[DocumentBlock(kind="paragraph", text=_clean_inline(markdown))]))
    return ast


def qa_document_ast(ast: DocumentAST, expected_sections: list[str] | None = None) -> list[DocumentQaIssue]:
    issues: list[DocumentQaIssue] = []
    if not ast.title.strip():
        issues.append(DocumentQaIssue("missing_title", "Document AST has no title."))
    if not ast.sections:
        issues.append(DocumentQaIssue("missing_sections", "Document AST has no sections."))
    empty = [section.heading for section in ast.sections if not section.blocks]
    if empty:
        issues.append(DocumentQaIssue("empty_sections", f"Sections without content: {', '.join(empty[:5])}."))
    for section in ast.sections:
        for block in section.blocks:
            if block.kind == "table":
                widths = {len(row) for row in block.rows}
                if len(widths) > 1:
                    issues.append(DocumentQaIssue("invalid_table", f"Table in '{section.heading}' has uneven rows."))
    if expected_sections:
        actual = {_canonical_heading(section.heading) for section in ast.sections}
        missing = [heading for heading in expected_sections if _canonical_heading(heading) not in actual]
        if missing:
            issues.append(DocumentQaIssue("missing_planned_sections", f"Missing planned sections: {', '.join(missing[:5])}."))
    return issues


def render_docx_from_markdown(title: str, markdown: str, *, expected_sections: list[str] | None = None) -> tuple[bytes, list[DocumentQaIssue]]:
    ast = document_ast_from_markdown(title, markdown)
    issues = qa_document_ast(ast, expected_sections=expected_sections)
    return render_docx_from_ast(ast), issues


def render_docx_from_ast(ast: DocumentAST) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True

    title = doc.add_heading(ast.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section_index, doc_section in enumerate(ast.sections):
        if section_index and doc_section.level <= 1:
            doc.add_section(WD_SECTION.CONTINUOUS)
        doc.add_heading(doc_section.heading, level=max(1, min(3, doc_section.level)))
        for block in doc_section.blocks:
            if block.kind == "paragraph":
                doc.add_paragraph(block.text)
            elif block.kind == "bullets":
                for item in block.items:
                    doc.add_paragraph(item, style="List Bullet")
            elif block.kind == "numbered_list":
                for item in block.items:
                    doc.add_paragraph(item, style="List Number")
            elif block.kind == "code":
                para = doc.add_paragraph()
                run = para.add_run(block.text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            elif block.kind == "table" and block.rows:
                _add_table(doc, block.rows)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_table(doc, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index].cells
        for col_index, value in enumerate(row):
            cells[col_index].text = value


def _is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and _is_table_separator(lines[index + 1].strip())


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    cells = [cell.strip() for cell in stripped.split("|")]
    return all(re.match(r"^:?-{3,}:?$", cell) for cell in cells if cell)


def _normalize_table_rows(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    return [row + [""] * (width - len(row)) for row in rows]


def _clean_heading(text: str) -> str:
    return _clean_inline(re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text).strip())


def _clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return _clean_text(text)


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _same_heading(left: str, right: str) -> bool:
    return _canonical_heading(left) == _canonical_heading(right)


def _canonical_heading(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()
