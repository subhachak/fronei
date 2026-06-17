from __future__ import annotations

import base64
import io
import logging
import re
import time
from dataclasses import dataclass
from typing import Any

import httpx

from app.config import get_settings
from app.services.agent_v3.models import Artifact, Source, ToolCall

logger = logging.getLogger(__name__)


def safe_filename(title: str, suffix: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "", title).strip().replace(" ", "-").lower()
    stem = re.sub(r"-{2,}", "-", stem)[:80].strip("-") or "agent-v3-output"
    return f"{stem}.{suffix.lstrip('.')}"


@dataclass
class AgentV3Tools:
    you_api_key: str | None = None
    tavily_api_key: str | None = None
    nimble_api_key: str | None = None
    nimble_api_endpoint: str = "https://api.webit.live/api/v1/realtime/serp"

    @classmethod
    def from_settings(cls) -> "AgentV3Tools":
        settings = get_settings()
        return cls(
            you_api_key=settings.you_api_key,
            tavily_api_key=settings.tavily_api_key,
            nimble_api_key=settings.nimble_api_key,
            nimble_api_endpoint=settings.nimble_api_endpoint,
        )

    def search_web(self, query: str, max_results: int = 6) -> tuple[list[Source], ToolCall]:
        started = time.perf_counter()
        tool = ToolCall(name="web_search", input={"query": query, "max_results": max_results})
        try:
            if not self.you_api_key and not self.tavily_api_key and not self.nimble_api_key:
                tool.ok = False
                tool.error = "YOU_API_KEY / TAVILY_API_KEY / NIMBLE_API_KEY is not configured"
                return [], tool

            errors: list[str] = []
            if self.you_api_key:
                try:
                    sources = self._search_you(query, max_results)
                    if sources:
                        tool.output = {"provider": "You.com", "source_count": len(sources)}
                        return sources, tool
                    errors.append("You.com returned no results")
                except Exception as exc:
                    logger.warning("agent_v3 You.com web_search failed: %s", exc)
                    errors.append(f"You.com: {exc}")

            if self.tavily_api_key:
                try:
                    response = httpx.post(
                        "https://api.tavily.com/search",
                        json={
                            "api_key": self.tavily_api_key,
                            "query": query,
                            "max_results": max_results,
                            "search_depth": "advanced",
                            "include_answer": False,
                            "include_raw_content": False,
                        },
                        timeout=20,
                    )
                    response.raise_for_status()
                    payload = response.json()
                    sources = [
                        Source(
                            title=str(item.get("title") or ""),
                            url=str(item.get("url") or ""),
                            snippet=str(item.get("content") or item.get("snippet") or ""),
                        )
                        for item in payload.get("results", [])
                        if isinstance(item, dict)
                    ]
                    if sources:
                        tool.output = {"provider": "Tavily", "source_count": len(sources)}
                        return sources, tool
                    errors.append("Tavily returned no results")
                except Exception as exc:
                    logger.warning("agent_v3 web_search failed: %s", exc)
                    errors.append(f"Tavily: {exc}")

            if self.nimble_api_key:
                try:
                    sources = self._search_nimble(query, max_results)
                    if sources:
                        tool.output = {"provider": "Nimble", "source_count": len(sources)}
                        return sources, tool
                    errors.append("Nimble returned no results")
                except Exception as exc:
                    logger.warning("agent_v3 Nimble web_search failed: %s", exc)
                    errors.append(f"Nimble: {exc}")

            tool.ok = False
            tool.error = "; ".join(errors) or "No web search provider returned results"
            tool.output = {"source_count": 0}
            return [], tool
        finally:
            tool.latency_ms = int((time.perf_counter() - started) * 1000)

    def _search_you(self, query: str, max_results: int) -> list[Source]:
        response = httpx.get(
            "https://api.ydc-index.io/search",
            headers={"X-API-Key": self.you_api_key or ""},
            params={"query": query, "num_web_results": max_results},
            timeout=20,
        )
        response.raise_for_status()
        payload = response.json()
        results = payload.get("hits") or payload.get("results") or payload.get("web_results") or []
        sources: list[Source] = []
        for item in results[:max_results]:
            if not isinstance(item, dict):
                continue
            url = str(item.get("url") or item.get("link") or "")
            snippets = item.get("snippets") or item.get("highlights") or []
            if isinstance(snippets, list):
                snippet = " ".join(str(part) for part in snippets)
            else:
                snippet = str(snippets or "")
            snippet = snippet or str(item.get("description") or item.get("snippet") or item.get("content") or "")
            if url:
                sources.append(
                    Source(
                        title=str(item.get("title") or ""),
                        url=url,
                        snippet=snippet,
                    )
                )
        return sources

    def _search_nimble(self, query: str, max_results: int) -> list[Source]:
        payload = {
            "query": query,
            "search_engine": "google_search",
            "country": "US",
            "locale": "en",
            "parse": True,
        }
        response = httpx.get(
            self.nimble_api_endpoint,
            headers={"Content-Type": "application/json", "Authorization": _nimble_auth_header(self.nimble_api_key or "")},
            params=payload,
            timeout=20,
        )
        if response.status_code == 405:
            response = httpx.post(
                self.nimble_api_endpoint,
                headers={"Content-Type": "application/json", "Authorization": _nimble_auth_header(self.nimble_api_key or "")},
                json=payload,
                timeout=20,
            )
        response.raise_for_status()
        data = response.json()
        sources: list[Source] = []
        for item in _nimble_result_items(data)[:max_results]:
            if not isinstance(item, dict):
                continue
            url = str(item.get("url") or item.get("link") or item.get("href") or "")
            snippets = item.get("snippets") or item.get("extra_snippets") or []
            if isinstance(snippets, list):
                snippet = " ".join(str(part) for part in snippets)
            else:
                snippet = str(snippets or "")
            snippet = snippet or str(item.get("description") or item.get("snippet") or item.get("content") or "")
            if url:
                sources.append(
                    Source(
                        title=str(item.get("title") or ""),
                        url=url,
                        snippet=snippet,
                    )
                )
        return sources

    def extract_urls(self, urls: list[str], max_chars_per_source: int = 2500) -> tuple[list[Source], ToolCall]:
        started = time.perf_counter()
        unique_urls = [url for index, url in enumerate(urls) if url and url not in urls[:index]][:6]
        tool = ToolCall(name="read_url", input={"urls": unique_urls})
        if not unique_urls or not self.tavily_api_key:
            tool.output = {"source_count": 0}
            tool.latency_ms = int((time.perf_counter() - started) * 1000)
            return [], tool

        try:
            response = httpx.post(
                "https://api.tavily.com/extract",
                json={"api_key": self.tavily_api_key, "urls": unique_urls},
                timeout=25,
            )
            response.raise_for_status()
            payload = response.json()
            results = payload.get("results") or []
            extracted: list[Source] = []
            for item in results:
                if not isinstance(item, dict):
                    continue
                content = str(item.get("raw_content") or item.get("content") or "")
                extracted.append(
                    Source(
                        title=str(item.get("title") or ""),
                        url=str(item.get("url") or ""),
                        content=content[:max_chars_per_source],
                    )
                )
            tool.output = {"source_count": len(extracted)}
            return extracted, tool
        except Exception as exc:
            logger.warning("agent_v3 read_url failed: %s", exc)
            tool.ok = False
            tool.error = str(exc)
            return [], tool
        finally:
            tool.latency_ms = int((time.perf_counter() - started) * 1000)

    def make_markdown_artifact(self, title: str, markdown: str) -> Artifact:
        encoded = base64.b64encode(markdown.encode("utf-8")).decode("ascii")
        return Artifact(
            kind="markdown",
            filename=safe_filename(title, "md"),
            mime_type="text/markdown",
            base64_data=encoded,
        )

    def make_docx_artifact(self, title: str, markdown: str) -> Artifact:
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(title, level=0)
            for line in markdown.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:].strip(), level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:].strip(), level=2)
                elif stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
                else:
                    doc.add_paragraph(stripped)
            buf = io.BytesIO()
            doc.save(buf)
            payload = buf.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            kind = "docx"
            filename = safe_filename(title, "docx")
        except Exception as exc:
            logger.warning("agent_v3 docx artifact failed; returning markdown: %s", exc)
            payload = markdown.encode("utf-8")
            mime = "text/markdown"
            kind = "markdown"
            filename = safe_filename(title, "md")
        return Artifact(
            kind=kind,  # type: ignore[arg-type]
            filename=filename,
            mime_type=mime,
            base64_data=base64.b64encode(payload).decode("ascii"),
        )


def source_context(sources: list[Source]) -> str:
    lines: list[str] = []
    for idx, source in enumerate(sources, start=1):
        body = source.content or source.snippet
        lines.append(f"[S{idx}] {source.title}\nURL: {source.url}\n{body[:2500]}")
    return "\n\n".join(lines)


def _nimble_auth_header(api_key: str) -> str:
    if api_key.lower().startswith(("bearer ", "basic ")):
        return api_key
    return f"Bearer {api_key}"


def _nimble_result_items(data: dict) -> list[dict]:
    parsing = data.get("parsing") if isinstance(data.get("parsing"), dict) else {}
    entities = parsing.get("entities") if isinstance(parsing.get("entities"), dict) else {}
    for key in ("SearchResult", "OrganicResult", "organic_results", "search_results", "results"):
        value = entities.get(key)
        if isinstance(value, list):
            return value
    for key in ("organic_results", "search_results", "results", "items"):
        value = data.get(key)
        if isinstance(value, list):
            return value
    return []
